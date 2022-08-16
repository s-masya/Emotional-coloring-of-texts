from docx import Document

from dostoevsky.tokenization import RegexTokenizer
from dostoevsky.models import FastTextSocialNetworkModel

doc = Document()
table = doc.add_table(rows=60, cols=6)
table.style = 'Table Grid'
row = table.rows[0]
row.cells[0].text = '№'
row.cells[1].text = 'Текст'
row.cells[2].text = 'Позитивная окр. (%)'
row.cells[3].text = 'Нейтральная окр. (%)'
row.cells[4].text = 'Негативная окр. (%)'
row.cells[5].text = 'Итоговая окр.'

for i in range(1, 30):
    doc_text = Document('D:/python/Дз на лето/small/' + str(i) + '.docx')
    text = []
    for paragraph in doc_text.paragraphs:
        text.append(paragraph.text)
    text = '\n'.join(text)

    tokenizer = RegexTokenizer()
    model = FastTextSocialNetworkModel(tokenizer=tokenizer)

    results = model.predict([text])
    for sentiment in results:
        positive = sentiment.get('positive')
        neutral = sentiment.get('neutral')
        negative = sentiment.get('negative')
        if positive is None:
            positive = 0
        if neutral is None:
            neutral = 0
        if negative is None:
            negative = 0

    if (neutral > negative) and (neutral > positive):
        emotional = 'Нейтральная'
    elif (positive > negative) and (positive > neutral):
        emotional = 'Позитивная'
    elif (negative > positive) and (negative > neutral):
        emotional = 'Негативная'
    else:
            emotional = '-'
    
    row = table.rows[i]
    row.cells[0].text = str(round(i, 5))
    row.cells[1].text = text
    row.cells[2].text = str(round(positive, 5))
    row.cells[3].text = str(round(neutral, 5))
    row.cells[4].text = str(round(negative, 5))
    row.cells[5].text = emotional

doc.save('D:\python\Дз на лето\Таблица результатов.docx')
